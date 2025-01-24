from django.shortcuts import render, redirect
from .forms import Inputform
from .models import Inputmodel
import joblib
import pandas as pd
from django.db.models import Count, Q, Avg, Sum
import json
from django.http import HttpResponse
from reportlab.lib.pagesizes import landscape, A1
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from xhtml2pdf import pisa
from io import BytesIO
from docx import Document


# Load the pre-trained model
preprocessing = joblib.load('savedmodels/prepro_pipeline.pkl')
model = joblib.load('savedmodels/model.pkl')

def home(request):
    return render(request, 'leadpredict/home.html')

def predict(request):
    if request.method == 'POST':
        form = Inputform(request.POST)
        if form.is_valid():
            data = {
                'TotalVisits': form.cleaned_data['total_visits'],
                'Total Time Spent on Website': form.cleaned_data['total_time_spent_on_website'],
                'Page Views Per Visit': form.cleaned_data['page_views_per_visit'],
                'Lead Origin': form.cleaned_data['lead_origin'],
                'Lead Source': form.cleaned_data['lead_source'],
                'Last Activity': form.cleaned_data['last_activity'],
                'Specialization': form.cleaned_data['specialization'],
                'What is your current occupation': form.cleaned_data['current_occupation'],
                'Tags': form.cleaned_data['tags'],
                'Lead Quality': form.cleaned_data['lead_quality'],
                'City': form.cleaned_data['city'],
                'Do Not Email': form.cleaned_data['do_not_email'],
                'Last Notable Activity': form.cleaned_data['last_notable_activity'],
                'A free copy of Mastering The Interview': form.cleaned_data['free_copy_of_mastering_the_interview'],
            }

            X = pd.DataFrame([data])
            try:
                # Preprocess and predict
                X_scaled = preprocessing.transform(X)
                stat = model.predict(X_scaled)[0]
                status = 'Converted' if stat == 1 else 'Not Converted'
                probability = model.predict_proba(X_scaled)[0][1]
                case = ' Hot case ' if probability > 0.7 else ' Potential ' if 0.3 < probability <= 0.7 else ' Cold case '

                # Save the form data including status, probability, and case
                inputmodel = Inputmodel(
                    total_visits=data['TotalVisits'],
                    total_time_spent_on_website=data['Total Time Spent on Website'],
                    page_views_per_visit=data['Page Views Per Visit'],
                    lead_origin=data['Lead Origin'],
                    lead_source=data['Lead Source'],
                    last_activity=data['Last Activity'],
                    specialization=data['Specialization'],
                    current_occupation=data['What is your current occupation'],
                    tags=data['Tags'],
                    lead_quality=data['Lead Quality'],
                    city=data['City'],
                    do_not_email=data['Do Not Email'],
                    last_notable_activity=data['Last Notable Activity'],
                    free_copy_of_mastering_the_interview=data['A free copy of Mastering The Interview'],
                    status=status,
                    probability=probability,
                    case=case
                )
                inputmodel.save()

                # Render the result page with the prediction details
                return render(request, 'leadpredict/result.html', {
                    'status': status,
                    'probability': probability,
                    'case': case,
                })

            except ValueError as e:
                form.add_error(None, f"Preprocessing error: {e}")

    else:
        form = Inputform()

    # Render the form with any validation errors if the method is GET or the form is invalid
    return render(request, 'leadpredict/inputform.html', {'form': form})

def display_table(request):
    # Order data by descending ID to show the most recent first
    input_data = Inputmodel.objects.all().order_by('-id')
    return render(request, 'leadpredict/display_table.html', {'input_data': input_data})

def download_pdf(request):
    data = Inputmodel.objects.all().values()
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="lead_data.pdf"'

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A1))
    elements = []

    table_data = []
    headers = list(data[0].keys())
    table_data.append(headers)

    for record in data:
        row = list(record.values())
        table_data.append(row)

    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(table)
    doc.build(elements)

    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)
    return response

def download_excel(request):
    data = Inputmodel.objects.all().values()
    df = pd.DataFrame(data)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="lead_data.xlsx"'

    with pd.ExcelWriter(response, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Lead Data')

    return response

def download_word(request):
    data = Inputmodel.objects.all().values()
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="lead_data.docx"'

    document = Document()
    document.add_heading('Lead Data', 0)

    table = document.add_table(rows=1, cols=len(data[0]))
    hdr_cells = table.rows[0].cells
    headers = list(data[0].keys())
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for record in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(record.values()):
            row_cells[i].text = str(value)

    document.save(response)
    return response

def download_json(request):
    data = list(Inputmodel.objects.all().values())
    response = HttpResponse(content_type='application/json')
    response['Content-Disposition'] = 'attachment; filename="lead_data.json"'
    response.write(json.dumps(data, indent=4))
    return response

def dashboard(request):
    try:
        lead_origin_p = Inputmodel.objects.values('lead_origin').annotate(total=Count('id'))
        lead_source_p = Inputmodel.objects.values('lead_source').annotate(total=Count('lead_source'))
        last_activity_p = Inputmodel.objects.values('last_activity').annotate(total=Count('last_activity'))
        specialization_p = Inputmodel.objects.values('specialization').annotate(total=Count('specialization'))
        lead_quality_p = Inputmodel.objects.values('lead_quality').annotate(total=Count('lead_quality'))
        city_p = Inputmodel.objects.values('city').annotate(total=Count('city'))
        total_visitors = Inputmodel.objects.count()
        avg_time_spent = Inputmodel.objects.aggregate(Avg('total_time_spent_on_website'))['total_time_spent_on_website__avg']
        avg_total_visits = Inputmodel.objects.aggregate(Avg('total_visits'))['total_visits__avg']
        converted_leads_count = Inputmodel.objects.filter(status='Converted').count()
        conversion_percentage = (converted_leads_count / total_visitors * 100) if total_visitors > 0 else 0
        avg_page_views_per_visit = Inputmodel.objects.aggregate(Avg('page_views_per_visit'))['page_views_per_visit__avg']

        lead_origin_data = Inputmodel.objects.values('lead_origin').annotate(
            converted_count=Count('id', filter=Q(status='Converted')),
            not_converted_count=Count('id', filter=Q(status='Not Converted'))
        )

        lead_source_data = Inputmodel.objects.values('lead_source').annotate(
            converted_count=Count('id', filter=Q(status='Converted')),
            not_converted_count=Count('id', filter=Q(status='Not Converted'))
        )
        last_activity_data = Inputmodel.objects.values('last_activity').annotate(
            converted_count=Count('id', filter=Q(status='Converted')),
            not_converted_count=Count('id', filter=Q(status='Not Converted'))
        )
        specialization_data = Inputmodel.objects.values('specialization').annotate(
            converted_count=Count('id', filter=Q(status='Converted')),
            not_converted_count=Count('id', filter=Q(status='Not Converted'))
        )
        lead_quality_data = Inputmodel.objects.values('lead_quality').annotate(
            converted_count=Count('id', filter=Q(status='Converted')),
            not_converted_count=Count('id', filter=Q(status='Not Converted'))
        )
        city_data = Inputmodel.objects.values('city').annotate(
            converted_count=Count('id', filter=Q(status='Converted')),
            not_converted_count=Count('id', filter=Q(status='Not Converted'))
        )
        
        context = {
            'total_visitors': total_visitors,
            'avg_time_spent': avg_time_spent,
            'avg_total_visits': avg_total_visits,
            'converted_leads_count': converted_leads_count,
            'conversion_percentage': conversion_percentage,
            'avg_page_views_per_visit': avg_page_views_per_visit,
            'lead_origin_p': lead_origin_p,
            'lead_source_p': lead_source_p,
            'last_activity_p': last_activity_p,
            'specialization_p': specialization_p,
            'lead_quality_p': lead_quality_p,
            'city_p': city_p,
            'lead_origin_data': lead_origin_data,
            'lead_source_data': lead_source_data,
            'last_activity_data': last_activity_data,
            'specialization_data': specialization_data,
            'lead_quality_data': lead_quality_data,
            'city_data': city_data,
        }

        return render(request, 'leadpredict/dashboard.html', context)

    except Exception as e:
        print(e)
        return render(request, 'leadpredict/dashboard.html', {'error': str(e)})
